import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
import re

# Fayllar
excel_file = "list.xlsx"
template_file = "template2.docx"
output_folder = "output"
logo_file = "logo.jpg"

# Excel oxu
df = pd.read_excel(excel_file, header=0, usecols="A:F")
df.columns = [
    "OTAĞIN ADI",
    "TƏHVİL ALDI",
    "İNVENTAR ADI",
    "İNVENTAR MARKASI",
    "İNVENTAR MODELİ",
    "İNVENTAR NÖMRƏSİ"
]
df.fillna("", inplace=True)

# Output qovluğu
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Təhlükəsiz fayl adı
def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", str(name))

# Hər şəxs üçün qruplaşdır
grouped = df.groupby("TƏHVİL ALDI")

for person, group in grouped:
    doc = Document(template_file)

    otagin_adi = str(group.iloc[0]["OTAĞIN ADI"]).strip() or "—"
    tehvil_aldi = str(person).strip() or "—"

    # Paragraph-ları yoxla
    for paragraph in doc.paragraphs:
        # Logo
        if "{{logo}}" in paragraph.text:
            paragraph.text = ""  # placeholder-ı silirik
            run = paragraph.add_run()
            run.add_picture(logo_file, width=Inches(1.4))  # logo ölçüsü

        # Title
        if "{{title}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{title}}", "TƏHVİL TƏSLİM АКТI \n(İnventarların təhvil-təslim edilməsinə dair)")

        # OTAĞIN ADI və TƏHVİL ALDI yuxarı paragraph-da gizlətmək
        if "{{otagin_adi}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{otagin_adi}}", otagin_adi)
            paragraph.runs[0].font.color.rgb = None  # rəng qalır görünməz etmək istəsən əlavə edə bilərsən

        if "{{tehvil_aldi}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{tehvil_aldi}}", tehvil_aldi)
            paragraph.runs[0].font.color.rgb = None

    # Table doldurma (ilk table)
    table = doc.tables[0]
    while len(table.rows) - 1 < len(group):
        table.add_row()

    for i, (_, row) in enumerate(group.iterrows()):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(row["İNVENTAR ADI"])
        row_cells[2].text = str(row["İNVENTAR MARKASI"])
        row_cells[3].text = str(row["İNVENTAR MODELİ"])
        row_cells[4].text = str(row["İNVENTAR NÖMRƏSİ"])

    # Word faylını yadda saxla
    word_file = os.path.join(output_folder, f"{sanitize_filename(person)}.docx")
    doc.save(word_file)

# PDF-ə çevir
convert(output_folder)

print("Hazırdır ✅ Logo, title və məlumatlar tam işlək şəkildə əlavə olundu.")